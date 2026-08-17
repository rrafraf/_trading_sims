from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "reports" / "human" / "M55_Vega_Coordination_Ledger.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "5B677A"
HEADER_FILL = "E8EEF5"
LIGHT_FILL = "F4F6F9"
GREEN_FILL = "EAF6EF"
GOLD_FILL = "FFF4CC"
RED_FILL = "FCE8E6"
GRID = "B9C3CF"


def set_run_font(run, *, size=None, bold=None, color=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=GRID):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_keep_with_next(paragraph, value=True):
    paragraph.paragraph_format.keep_with_next = value


def paragraph(doc, text="", style=None, *, bold_prefix=None, color=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True, color=color)
        r2 = p.add_run(text[len(bold_prefix) :])
        set_run_font(r2, color=color)
    else:
        r = p.add_run(text)
        set_run_font(r, color=color)
    return p


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    set_keep_with_next(p)
    for run in p.runs:
        set_run_font(run, bold=True, color=BLUE if level < 3 else DARK_BLUE)
    return p


def write_cell(cell, text, *, bold=False, color=None, align=None):
    p = cell.paragraphs[0]
    p.alignment = align or WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_run_font(r, bold=bold, color=color)


def add_table(doc, headers, rows, widths, status_col=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    set_table_borders(table)

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, HEADER_FILL)
        write_cell(cell, header, bold=True, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER)

    for row_values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_values):
            if status_col is not None and i == status_col:
                status = str(value).upper()
                fill = LIGHT_FILL
                if "WORKING" in status or "DONE" in status:
                    fill = GREEN_FILL
                elif "WAITING" in status:
                    fill = GOLD_FILL
                elif "BLOCKED" in status:
                    fill = RED_FILL
                set_cell_shading(cells[i], fill)
                write_cell(cells[i], str(value), bold=True, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER)
            else:
                write_cell(cells[i], str(value))

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    return table


def configure_doc(doc):
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for attr in ("top_margin", "right_margin", "bottom_margin", "left_margin"):
        setattr(section, attr, Inches(1.0))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("M55 + Vega coordination ledger")
    set_run_font(r, size=9, color=MUTED)


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    configure_doc(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("M55 + Vega Coordination Ledger")
    set_run_font(run, size=24, bold=True, color=INK)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(8)
    r = subtitle.add_run("Human-facing checkpoint for Rafa. Last updated: 2026-08-14.")
    set_run_font(r, size=10.5, color=MUTED)

    paragraph(
        doc,
        "Purpose: Keep the machine-speed work legible. This document records what each agent is doing next, what is waiting on Rafa, and which decisions are already settled.",
        bold_prefix="Purpose:",
    )

    add_heading(doc, "Current Snapshot", 1)
    add_table(
        doc,
        ["Area", "State", "Meaning for Rafa"],
        [
            ["Bridge direction", "DECIDED", "Do not merge repos yet. Use a file/subprocess bridge first."],
            ["_trading_sims", "WORKING SURFACE", "Owns benchmark harness, adapters, canonical scenarios, comparison reports."],
            ["muni", "ENGINE SURFACE", "Owns execution semantics, trace/evidence UI, anomaly vocabulary."],
            ["Implementation", "WAITING ON RAFA", "No code bridge starts until Rafa approves the next bounded chunk."],
            ["Credit discipline", "ACTIVE RULE", "No new agent task, fork, or recurring watcher without explicit approval."],
        ],
        [1875, 1500, 5985],
        status_col=1,
    )

    add_heading(doc, "Agent Status", 1)
    add_table(
        doc,
        ["Agent", "Status", "Next Thing", "Stop Condition"],
        [
            [
                "M55",
                "WAITING ON RAFA",
                "If approved, implement the muni target-position-v1 CLI seam. Otherwise maintain this ledger and coordinate decisions only.",
                "Stops after a tested bounded chunk, or when a human decision is required.",
            ],
            [
                "Vega",
                "WAITING",
                "Benchmark harness and bridge contract are captured. Can answer what M55 should reuse first, or later add muni_adapter.py after the muni CLI exists.",
                "Stops until Rafa or M55 asks for a specific benchmark-side move.",
            ],
            [
                "Rafa",
                "HUMAN OWNER",
                "Choose the next bounded move when ready: implement CLI seam, ask Vega for reuse list, or pause coding.",
                "Final authority. Agent-to-agent notes are advisory only.",
            ],
        ],
        [1180, 1450, 4480, 2250],
        status_col=1,
    )

    add_heading(doc, "Shared Rules", 1)
    add_table(
        doc,
        ["Rule", "Text"],
        [
            ["Visible stop state", "Every agent stop should say DONE, WORKING UNTIL X, WAITING ON RAFA, WAITING ON VEGA, or BLOCKED."],
            ["Authority", "Agents may recommend. Rafa decides project direction and credit-spending actions."],
            ["Bridge first", "First bridge mode is target-position-v1 for cheap parity. Second mode is order-intent-v1 for hard execution semantics."],
            ["No silent waiting", "If an agent pauses, it must say what it is waiting on in the task and update this ledger when practical."],
        ],
        [1875, 7485],
    )

    add_heading(doc, "Next Choices", 1)
    add_table(
        doc,
        ["Choice", "Cost", "Result"],
        [
            ["Approve M55 CLI seam", "Medium", "Creates the real muni output contract needed for _trading_sims to compare muni as an engine."],
            ["Ask Vega for reuse list", "Low", "Produces a benchmark-side shortlist of code/tests/contracts to reuse before M55 edits muni."],
            ["Pause coding", "Low", "Keeps current state stable. Agents only update status when you return."],
        ],
        [2400, 1200, 5760],
    )

    add_heading(doc, "Update Protocol For Agents", 1)
    paragraph(
        doc,
        "When M55 or Vega makes meaningful progress, update the Agent Status row and Current Snapshot if the human-facing state changed. Keep entries short. Do not turn this into a full engineering journal.",
    )
    paragraph(
        doc,
        "If a change would spend credits, create/fork a task, start recurring work, or change repo direction, mark WAITING ON RAFA instead of proceeding silently.",
    )

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build())
